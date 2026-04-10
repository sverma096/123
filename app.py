from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4, legal
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import os
import re
import uuid

app = Flask(__name__)

OUTPUT_DIR = "generated"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Optional OpenAI support
OPENAI_ENABLED = False
try:
    from openai import OpenAI
    if os.getenv("OPENAI_API_KEY"):
        client = OpenAI()
        OPENAI_ENABLED = True
except Exception:
    OPENAI_ENABLED = False


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    replacements = {
        r"\bfull stop\b": ".",
        r"\bcomma\b": ",",
        r"\bcolon\b": ":",
        r"\bsemicolon\b": ";",
        r"\bquestion mark\b": "?",
        r"\bexclamation mark\b": "!",
        r"\bnew paragraph\b": "\n\n",
        r"\bnext line\b": "\n",
        r"पूर्ण विराम": "।",
        r"कॉमा": ",",
        r"नया पैराग्राफ": "\n\n",
        r"नई लाइन": "\n",
    }

    for pattern, value in replacements.items():
        text = re.sub(pattern, value, text, flags=re.IGNORECASE)

    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    # Simple capitalization for Latin script
    parts = re.split(r'(?<=[.!?।])\s+', text)
    parts = [p[:1].upper() + p[1:] if p else p for p in parts]
    return " ".join(parts).strip()


def local_ai_cleanup(text: str, language_mode: str, doc_type: str) -> str:
    text = normalize_text(text)

    if doc_type == "notice":
        prefix = "LEGAL NOTICE\n\n"
    elif doc_type == "agreement":
        prefix = "AGREEMENT\n\n"
    elif doc_type == "reply":
        prefix = "LEGAL REPLY\n\n"
    else:
        prefix = ""

    return prefix + text


def openai_cleanup(text: str, language_mode: str, doc_type: str) -> str:
    if not OPENAI_ENABLED:
        return local_ai_cleanup(text, language_mode, doc_type)

    system_prompt = (
        "You are a legal drafting assistant. "
        "Clean the user's dictated text, fix punctuation, paragraphing, and grammar, "
        "preserve meaning, and format it professionally. "
        "Do not invent facts. "
        "If the user text is mixed Hindi and English, preserve the mixed language naturally."
    )

    user_prompt = f"""
Language mode: {language_mode}
Document type: {doc_type}

Please clean and format this dictated legal draft:

{text}
"""

    try:
        response = client.responses.create(
            model="gpt-5",
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        return response.output_text.strip()
    except Exception:
        return local_ai_cleanup(text, language_mode, doc_type)


def build_document(title: str, body: str, include_signature: bool, include_stamp: bool) -> str:
    title = (title or "LEGAL DOCUMENT").strip().upper()

    lines = [title, "", body.strip(), "", "Place: ____________________", "Date: _____________________", ""]

    if include_signature:
        lines.extend([
            "Signature: __________________________",
            "Authorized Signatory",
            ""
        ])

    if include_stamp:
        lines.append("[STAMP HERE]")

    return "\n".join(lines).strip()


@app.route("/")
def home():
    return render_template("index.html", openai_enabled=OPENAI_ENABLED)


@app.post("/preview")
def preview():
    data = request.get_json(force=True)

    title = data.get("title", "Legal Document")
    text = data.get("text", "")
    language_mode = data.get("language_mode", "hinglish")
    doc_type = data.get("doc_type", "general")
    use_ai = bool(data.get("use_ai", True))
    include_signature = bool(data.get("signature", True))
    include_stamp = bool(data.get("stamp", True))

    cleaned = openai_cleanup(text, language_mode, doc_type) if use_ai else normalize_text(text)
    final_text = build_document(title, cleaned, include_signature, include_stamp)

    return jsonify({
        "preview": final_text,
        "openai_enabled": OPENAI_ENABLED
    })


@app.post("/generate")
def generate():
    title = request.form.get("title", "Legal Document")
    text = request.form.get("text", "")
    format_type = request.form.get("format", "legal")
    file_type = request.form.get("filetype", "word")
    language_mode = request.form.get("language_mode", "hinglish")
    doc_type = request.form.get("doc_type", "general")

    use_ai = request.form.get("use_ai") == "on"
    include_signature = request.form.get("signature") == "on"
    include_stamp = request.form.get("stamp") == "on"

    cleaned = openai_cleanup(text, language_mode, doc_type) if use_ai else normalize_text(text)
    final_text = build_document(title, cleaned, include_signature, include_stamp)

    uid = uuid.uuid4().hex[:8]

    if file_type == "word":
        filename = f"lexvoice_{uid}.docx"
        path = os.path.join(OUTPUT_DIR, filename)

        doc = Document()
        section = doc.sections[0]

        if format_type == "legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)
        else:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.strip().upper())
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph("")

        first_line_skipped = False
        for line in final_text.split("\n"):
            if not first_line_skipped and line.strip().upper() == title.strip().upper():
                first_line_skipped = True
                continue
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(8)
            para.add_run(line)

        doc.save(path)

    else:
        filename = f"lexvoice_{uid}.pdf"
        path = os.path.join(OUTPUT_DIR, filename)
        page_size = legal if format_type == "legal" else A4

        pdf = SimpleDocTemplate(
            path,
            pagesize=page_size,
            leftMargin=72,
            rightMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "LegalTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=16
        )

        body_style = ParagraphStyle(
            "LegalBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=11,
            leading=17,
            alignment=TA_LEFT,
            spaceAfter=10
        )

        story = [Paragraph(title.strip().upper(), title_style), Spacer(1, 4)]

        first_line_skipped = False
        for line in final_text.split("\n"):
            if not first_line_skipped and line.strip().upper() == title.strip().upper():
                first_line_skipped = True
                continue

            if line.strip():
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))
            else:
                story.append(Spacer(1, 8))

        pdf.build(story)

    return send_file(path, as_attachment=True, download_name=filename)


if __name__ == "__main__":
    app.run(debug=True)
